"""
Post 9 — Generate 3 versions of the LinkedIn text as Word doc.

Structure per version:
  - Hook (opening)
  - 3-4 key findings (bullets)
  - Standout single-line insight
  - Method one-liner
  - Call-to-action / question

Versions:
  A. English (global reach)
  B. Português (BR/PT direto)
  C. Hybrid (EN main + PT punchline)

Output: Outputs/post9/post9_linkedin_texts.docx
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

_here = Path(__file__).resolve().parent
cowork_root = _here.parent.parent
OUT_DIR = cowork_root / "Outputs" / "post9"
OUT_DIR.mkdir(parents=True, exist_ok=True)
OUT_FILE = OUT_DIR / "post9_linkedin_texts.docx"

doc = Document()

# Style defaults
style_normal = doc.styles["Normal"]
style_normal.font.name = "Arial"
style_normal.font.size = Pt(11)

for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)


def add_title(text, size=18, color=(0, 0, 0)):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(size)
    r.font.bold = True
    r.font.color.rgb = RGBColor(*color)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)


def add_subheading(text, color=(90, 90, 90)):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(10)
    r.font.italic = True
    r.font.color.rgb = RGBColor(*color)
    p.paragraph_format.space_after = Pt(4)


def add_para(text, size=11, bold=False, italic=False, color=(0, 0, 0)):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    r.font.color.rgb = RGBColor(*color)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.runs[0] if p.runs else p.add_run("")
    r.text = text
    r.font.name = "Arial"
    r.font.size = Pt(11)


def add_divider():
    p = doc.add_paragraph()
    r = p.add_run("─" * 60)
    r.font.color.rgb = RGBColor(180, 180, 180)
    r.font.size = Pt(9)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(10)


# =============================================================================
# COVER
# =============================================================================

add_title("Post 9 — Container carrier deployment, Brazilian ports 2025", size=16)
add_subheading("Three versions for LinkedIn — EN, PT, Hybrid  •  choose one, edit if needed")

add_para(
    "Base numbers used across all versions:", bold=True
)
add_para(
    "4.226 container vessel calls  •  433 unique IMOs  •  9 Brazilian ports  •  "
    "15 container terminals  •  74,5% carrier-identified  •  Source: ANTAQ 2025 + VesselFinder",
    size=10, italic=True, color=(90, 90, 90)
)

add_divider()

# =============================================================================
# VERSION A — ENGLISH
# =============================================================================

add_title("A — English  (global reach, Trade Lane Manager tone)")

add_para(
    "Nine Brazilian container ports. Two very different worlds for procurement.",
    bold=True, size=12
)
add_para("")

add_para(
    "I pulled every 2025 international container call from ANTAQ, cross-checked "
    "vessel names on VesselFinder, and grouped carriers by parent and by alliance "
    "(Feb 2025 config). The Herfindahl index tells the story cleanly:"
)
add_para("")

add_para("Five ports — highly concentrated  (HHI ≥ 2500)", bold=True)
add_bullet("Manaus 3.707  •  Itajaí 3.390  •  Pecém 2.970  •  Rio Grande 2.607  •  Salvador 2.564")
add_bullet("These are single-carrier or duopoly markets. MSC alone runs 40-48% of calls in four of them.")
add_para("")

add_para("Four ports — moderately concentrated  (HHI < 2300)", bold=True)
add_bullet("Rio de Janeiro 2.292  •  Santos 1.983  •  Paranaguá 1.895  •  Portonave 1.726")
add_bullet("Real optionality here. Santos alone has 8 carriers above 2% share.")
add_para("")

add_para("Three findings a procurement model tends to miss:", bold=True)
add_bullet("MSC is #1 in 9 of 9 ports — no exception. That is not competition, that is a plateau.")
add_bullet("Gemini Cooperation captured 21,9% of calls in five months (Feb-Jul 2025). Fast materialisation of the 2025 realignment.")
add_bullet("Portonave is the only port where Premier Alliance (ONE / HMM / YM) shows real presence. Everywhere else, ONE is a footnote.")
add_para("")

add_para(
    "One number worth keeping: the alliance-level HHI on identified carriers is 2.875. "
    "That is textbook \"highly concentrated\". Which means the negotiation surface "
    "for a Brazil-China Trade Lane Manager is not 12 carriers — it is 3 alliance blocks "
    "plus MSC standalone.",
    italic=True
)
add_para("")

add_para(
    "Method note: 74,5% of calls identified via vessel-name pattern matching + "
    "known-fleet enrichment. Non-identified are mostly charter tonnage without a "
    "carrier brand in the name. HHI reported both with and without UNKNOWN — "
    "conclusions hold under both. Full-year 2025 dataset used because ANTAQ 2026 "
    "is not yet consolidated; a 2026 refresh is planned once the data lands.",
    size=10, color=(90, 90, 90)
)
add_para("")

add_para(
    "Full data, code and methodology on GitHub → [link]",
    size=10, color=(90, 90, 90)
)
add_para("")
add_para(
    "Question for the shipping side: what would you look at that this analysis "
    "is not capturing?",
    bold=True
)

add_divider()

# =============================================================================
# VERSION B — PORTUGUÊS
# =============================================================================

add_title("B — Português  (direto ao ponto, tom BR/PT)")

add_para(
    "Nove portos brasileiros de contentor. Dois mundos completamente diferentes "
    "para quem faz procurement.",
    bold=True, size=12
)
add_para("")

add_para(
    "Puxei todas as atracações de longo curso container da ANTAQ 2025, "
    "cruzei com nomes de navios da VesselFinder, e agrupei por armador, "
    "por grupo e por aliança (config Feb 2025). O índice de Herfindahl "
    "(HHI) mostra:"
)
add_para("")

add_para("Cinco portos altamente concentrados  (HHI ≥ 2500)", bold=True)
add_bullet("Manaus 3.707  •  Itajaí 3.390  •  Pecém 2.970  •  Rio Grande 2.607  •  Salvador 2.564")
add_bullet("Aqui o mercado é de um armador só ou de duopólio. MSC sozinha faz 40-48% das escalas em quatro destes portos.")
add_para("")

add_para("Quatro portos moderadamente concentrados  (HHI < 2300)", bold=True)
add_bullet("Rio de Janeiro 2.292  •  Santos 1.983  •  Paranaguá 1.895  •  Portonave 1.726")
add_bullet("Aqui há alternativas. Santos tem 8 armadores acima de 2% de share.")
add_para("")

add_para("Três leituras que um modelo padrão de procurement costuma ignorar:", bold=True)
add_bullet("MSC é #1 em 9 de 9 portos. Não é concorrência — é dominância estrutural.")
add_bullet("A Gemini Cooperation ganhou 21,9% do mercado aliança-level em 5 meses (Feb-Jul 2025). A reconfiguração de 2025 materializou-se rápido no Brasil.")
add_bullet("Portonave é o único porto onde a Premier Alliance (ONE / HMM / YM) tem presença real. Nos outros portos, é resíduo.")
add_para("")

add_para(
    "Um número que vale reter: HHI aliança-level nos armadores identificados = 2.875. "
    "\"Altamente concentrado\" pela regra do DOJ/FTC. Ou seja, quem faz Trade Lane "
    "China-Brasil não está a negociar com 12 armadores — está a negociar com 3 "
    "blocos de aliança e a MSC em separado.",
    italic=True
)
add_para("")

add_para(
    "Nota metodológica: 74,5% das escalas identificadas via pattern matching de "
    "nome de navio + enriquecimento com base de frotas conhecidas. Os não-identificados "
    "são sobretudo charter sem marca de armador no nome. HHI reportado com e sem "
    "UNKNOWN — as conclusões aguentam nos dois cenários. Uso o dataset completo de "
    "2025 porque os dados ANTAQ 2026 ainda não estão consolidados; um refresh 2026 "
    "está previsto quando os dados sairem.",
    size=10, color=(90, 90, 90)
)
add_para("")

add_para(
    "Dados completos, código e methodology no GitHub → [link]",
    size=10, color=(90, 90, 90)
)
add_para("")
add_para(
    "Pergunta para quem trabalha shipping: o que é que esta análise está a deixar de fora?",
    bold=True
)

add_divider()

# =============================================================================
# VERSION C — HYBRID
# =============================================================================

add_title("C — Hybrid  (EN core + PT punchline, targeting BR-based shipping lines with global HQ)")

add_para(
    "Nine Brazilian container ports. Two very different procurement markets.",
    bold=True, size=12
)
add_para("")

add_para(
    "Herfindahl-Hirschman by port, all 2025 international container calls "
    "(ANTAQ + VesselFinder):"
)
add_para("")

add_para("Highly concentrated  →  Manaus, Itajaí, Pecém, Rio Grande, Salvador", bold=True)
add_bullet("HHI 2.564 – 3.707. Single-carrier markets. MSC 40-48% share in most.")
add_para("")

add_para("Moderately concentrated  →  Rio de Janeiro, Santos, Paranaguá, Portonave", bold=True)
add_bullet("HHI 1.726 – 2.292. Real optionality. Santos: 8 carriers >2%.")
add_para("")

add_para("What the numbers say — three findings that matter for procurement:", bold=True)
add_bullet("MSC leads all 9 ports (25-48% share). Not competition. Structural dominance.")
add_bullet("Gemini Cooperation grabbed 21,9% alliance-level share in 5 months post-launch.")
add_bullet("Portonave is the only port where the Premier Alliance is materially visible. Elsewhere, residual.")
add_para("")

add_para(
    "Alliance-level HHI (identified carriers) = 2.875. \"Highly concentrated\" by DOJ/FTC rule of thumb. "
    "The negotiation is not 12 carriers, it's 3 alliance blocks plus MSC standalone.",
    italic=True
)
add_para("")

add_para(
    "Um dado extra que salta ao olho: Portonave é a excepção brasileira. Único porto onde as 4 alianças "
    "têm presença significativa em simultâneo. Alguém aí do lado operacional consegue confirmar se isto "
    "reflete o mix de serviços feeders ou algo mais estrutural?",
    bold=True
)
add_para("")

add_para(
    "Method: 74,5% of calls identified via vessel-name pattern matching + known-fleet enrichment. "
    "Non-identified = mostly charter tonnage without carrier brand. HHI reported with and without "
    "UNKNOWN — conclusions robust. Full-year 2025 used because ANTAQ 2026 is not yet consolidated; "
    "a 2026 refresh is planned.",
    size=10, color=(90, 90, 90)
)
add_para("")
add_para(
    "Full analysis, code, chart, methodology → GitHub [link]",
    size=10, color=(90, 90, 90)
)

add_divider()

# =============================================================================
# NOTES
# =============================================================================

add_title("Editorial notes for Hugo", size=13)
add_bullet("All 3 versions open with the same hook: 9 ports split into 2 concentration groups.")
add_bullet("Version A is the safest for global shipping-line reach.")
add_bullet("Version B is more direct — good if you want to test the BR/PT market specifically.")
add_bullet("Version C ends with a question in PT — invites Brazilian ops folks to comment. Highest engagement bet.")
add_bullet("Hero number is HHI alliance = 2.875 — repeat it in comments if the post gains traction.")
add_bullet("Replace [link] with the real GitHub URL once the repo commit lands.")

doc.save(OUT_FILE)
print(f"Saved: {OUT_FILE}")
